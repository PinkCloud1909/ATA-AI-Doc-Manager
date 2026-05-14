import io
import logging
from html.parser import HTMLParser
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Mapping of content types and file extensions to extractor methods
_PDF_TYPES = {"application/pdf"}
_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_TEXT_TYPES = {"text/plain", "text/csv", "text/markdown"}
_HTML_TYPES = {"text/html"}

_PDF_EXTENSIONS = {".pdf"}
_DOCX_EXTENSIONS = {".docx"}
_TEXT_EXTENSIONS = {".txt", ".csv", ".md", ".log"}
_HTML_EXTENSIONS = {".html", ".htm"}


class _HTMLTextExtractor(HTMLParser):
    """Minimal HTML parser that strips tags and collects visible text."""

    _SKIP_TAGS = {"script", "style", "head", "noscript"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth: int = 0

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0 and data.strip():
            self._parts.append(data.strip())

    def get_text(self) -> str:
        return "\n".join(self._parts)


class TextExtractor:
    """Extracts plain text content from binary file data (PDF, DOCX, HTML, TXT)."""

    def extract(
        self,
        file_bytes: bytes,
        content_type: str | None = None,
        filename: str = "",
    ) -> str:
        """Dispatch to the appropriate extraction method based on content type
        or file extension.

        Returns the extracted plain text.
        Raises ``ValueError`` if the file type is not supported.
        """
        ct = (content_type or "").lower().strip()
        ext = Path(filename).suffix.lower() if filename else ""

        if ct in _PDF_TYPES or ext in _PDF_EXTENSIONS:
            return self._extract_pdf(file_bytes)
        if ct in _DOCX_TYPES or ext in _DOCX_EXTENSIONS:
            return self._extract_docx(file_bytes)
        if ct in _HTML_TYPES or ext in _HTML_EXTENSIONS:
            return self._extract_html(file_bytes)
        if ct in _TEXT_TYPES or ext in _TEXT_EXTENSIONS:
            return self._extract_text(file_bytes)

        raise ValueError(
            f"Unsupported file type: content_type={content_type!r}, "
            f"filename={filename!r}"
        )

    # --- private extractors ---

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from all pages of a PDF using PyMuPDF."""
        pages: list[str] = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text = page.get_text()
                if text:
                    pages.append(text)
        result = "\n".join(pages)
        logger.debug(
            "pdf_text_extracted",
            extra={"char_count": len(result), "page_count": len(pages)},
        )
        if not result.strip():
            logger.warning(
                "pdf_text_empty",
                extra={
                    "note": (
                        "PDF appears to be image-only or scanned. "
                        "OCR is required to extract text."
                    )
                },
            )
        return result

    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from paragraphs and table cells of a DOCX file."""
        doc = DocxDocument(io.BytesIO(file_bytes))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Include table cell content — often omitted but important for structured docs
        table_cells = [
            cell.text.strip()
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]

        result = "\n".join(paragraphs + table_cells)
        logger.debug(
            "docx_text_extracted",
            extra={
                "char_count": len(result),
                "paragraph_count": len(paragraphs),
                "table_cell_count": len(table_cells),
            },
        )
        return result

    def _extract_html(self, file_bytes: bytes) -> str:
        """Strip HTML tags and return visible text content."""
        raw = file_bytes.decode("utf-8", errors="replace")
        parser = _HTMLTextExtractor()
        parser.feed(raw)
        result = parser.get_text()
        logger.debug(
            "html_text_extracted",
            extra={"char_count": len(result)},
        )
        return result

    def _extract_text(self, file_bytes: bytes) -> str:
        """Decode raw bytes as UTF-8 text (plain text, CSV, Markdown, log)."""
        return file_bytes.decode("utf-8", errors="replace")
