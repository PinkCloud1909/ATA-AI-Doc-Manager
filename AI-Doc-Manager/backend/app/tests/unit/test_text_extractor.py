import io
import pytest

from app.modules.vectorization.domain.text_extractor import TextExtractor


@pytest.fixture
def extractor():
    return TextExtractor()


class TestExtractText:
    def test_plain_text_by_content_type(self, extractor):
        data = "Hello, world!\nThis is a test document."
        result = extractor.extract(
            data.encode("utf-8"), content_type="text/plain", filename="doc.txt"
        )
        assert "Hello, world!" in result
        assert "test document" in result

    def test_plain_text_by_extension(self, extractor):
        data = "Line one\nLine two"
        result = extractor.extract(
            data.encode("utf-8"), content_type=None, filename="notes.txt"
        )
        assert result == data

    def test_plain_text_markdown(self, extractor):
        data = "# Heading\n\nParagraph."
        result = extractor.extract(
            data.encode("utf-8"), content_type="text/markdown", filename="readme.md"
        )
        assert "# Heading" in result


class TestExtractPDF:
    def test_pdf_extraction(self, extractor):
        """Create a minimal PDF in-memory using PyMuPDF and extract its text."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from PDF")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extractor.extract(
            pdf_bytes, content_type="application/pdf", filename="test.pdf"
        )
        assert "Hello from PDF" in result

    def test_pdf_by_extension(self, extractor):
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Extension test")
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extractor.extract(pdf_bytes, content_type=None, filename="test.pdf")
        assert "Extension test" in result


class TestExtractDocx:
    def test_docx_extraction(self, extractor):
        """Create a minimal DOCX in-memory using python-docx and extract text."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph from DOCX")
        doc.add_paragraph("Second paragraph")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = extractor.extract(
            docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="test.docx",
        )
        assert "First paragraph from DOCX" in result
        assert "Second paragraph" in result

    def test_docx_by_extension(self, extractor):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Extension DOCX")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = extractor.extract(docx_bytes, content_type=None, filename="file.docx")
        assert "Extension DOCX" in result


class TestUnsupportedType:
    def test_unsupported_content_type_raises(self, extractor):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extractor.extract(b"data", content_type="image/png", filename="img.png")

    def test_no_type_info_raises(self, extractor):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extractor.extract(b"data", content_type=None, filename="")
